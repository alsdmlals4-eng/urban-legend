#!/usr/bin/env python3
"""Build and verify a deterministic DOCX mirror of the living Markdown GDD."""
from __future__ import annotations

import argparse
import hashlib
import io
import re
import sys
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "GAME_DESIGN_DOCUMENT.md"
OUTPUT = ROOT / "docs" / "URBAN_LEGEND_GAME_DESIGN.docx"
FORMAT_VERSION = "urban-legend-gdd-index-v3"
HASH_PREFIX = "gdd-source-sha256="
FONT = "NanumSquare"
NAVY = "1F2A44"
BLUE = "2E75B6"
TEXT = "222222"
MUTED = "666666"
LIGHT = "F3F5F8"


def args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--build", action="store_true")
    mode.add_argument("--check", action="store_true")
    return parser.parse_args()


def image_paths(markdown: str) -> list[Path]:
    return [(SOURCE.parent / raw).resolve() for raw in re.findall(r"!\[[^\]]*\]\(([^)]+)\)", markdown)]


def source_hash(markdown: str) -> str:
    digest = hashlib.sha256(FORMAT_VERSION.encode())
    digest.update(markdown.encode())
    for path in image_paths(markdown):
        digest.update(path.relative_to(ROOT).as_posix().encode())
        digest.update(path.read_bytes())
    return digest.hexdigest()


def clean(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("**", "").replace("`", "").strip()


def font(run, size: float = 9, color: str = TEXT, bold: bool = False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=width)
    table.autofit = True
    for row_index, values in enumerate(rows):
        for column, value in enumerate(values):
            cell = table.cell(row_index, column)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, NAVY if row_index == 0 else (LIGHT if row_index % 2 else "FFFFFF"))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(clean(value))
            font(run, 7.8, "FFFFFF" if row_index == 0 else TEXT, row_index == 0)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            table.rows[0]._tr.get_or_add_trPr().append(repeat)


def picture_stream(path: Path) -> io.BytesIO:
    with Image.open(path) as source:
        image = source.convert("RGB")
        image.thumbnail((1100, 850), Image.Resampling.LANCZOS)
        stream = io.BytesIO()
        image.save(stream, "JPEG", quality=72, optimize=True, progressive=True)
        stream.seek(0)
        return stream


def add_picture(doc: Document, path: Path, alt: str) -> None:
    with Image.open(path) as source:
        vertical = source.height > source.width * 1.25
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(picture_stream(path), width=Inches(4.0 if vertical else 6.7))
    shape._inline.docPr.set("title", alt)
    shape._inline.docPr.set("descr", alt)


def configure(doc: Document, digest: str) -> None:
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.72)
    section.top_margin = section.bottom_margin = Inches(0.65)
    doc.core_properties.title = "괴이 기록국 게임기획서"
    doc.core_properties.subject = "MVP-043 + CORE-VALIDATION-001 + UX-PD-001 2A / Ver 4.2"
    doc.core_properties.keywords = HASH_PREFIX + digest
    for style_name, size, color in (("Normal", 9, TEXT), ("Heading 1", 16, NAVY), ("Heading 2", 12, BLUE)):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("괴이 기록국 · GDD v1.3 · UX-PD-001 2A · Ver 4.2"), 8, MUTED)


def render(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("```"):
            block: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                block.append(lines[index]); index += 1
            table = doc.add_table(rows=1, cols=1)
            shade(table.cell(0, 0), LIGHT)
            run = table.cell(0, 0).paragraphs[0].add_run("\n".join(block))
            font(run, 8, TEXT)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1].strip(" |")):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                parts = [part.strip() for part in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", part.replace(" ", "")) for part in parts):
                    rows.append(parts)
                index += 1
            add_table(doc, rows)
            continue
        match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if match:
            add_picture(doc, (SOURCE.parent / match.group(2)).resolve(), match.group(1)); index += 1; continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(); paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(clean(line[2:])); font(run, 28, NAVY, True); index += 1; continue
        if line.startswith("## "):
            doc.add_paragraph(clean(line[3:]), style="Heading 1"); index += 1; continue
        if line.startswith("### "):
            doc.add_paragraph(clean(line[4:]), style="Heading 2"); index += 1; continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph(); paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(clean(line[2:])); font(run, 9, BLUE, True); index += 1; continue
        if line.startswith("- "):
            doc.add_paragraph(clean(line[2:]), style="List Bullet"); index += 1; continue
        if re.match(r"^\d+\. ", line):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", line)), style="List Number"); index += 1; continue
        paragraph_lines = [line]; index += 1
        while index < len(lines) and lines[index].strip() and not re.match(r"^(#|>|- |\d+\. |```|\||!\[)", lines[index].strip()):
            paragraph_lines.append(lines[index].strip()); index += 1
        run = doc.add_paragraph().add_run(clean(" ".join(paragraph_lines))); font(run)


def build() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    for path in image_paths(markdown):
        if not path.is_file():
            raise FileNotFoundError(path)
    digest = source_hash(markdown)
    document = Document()
    configure(document, digest)
    render(document, markdown)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"BUILT {OUTPUT}\nSOURCE_HASH {digest}")


def check() -> None:
    if not OUTPUT.is_file():
        raise FileNotFoundError(f"Missing generated DOCX: {OUTPUT}; run --build first")
    markdown = SOURCE.read_text(encoding="utf-8")
    expected = source_hash(markdown)
    document = Document(OUTPUT)
    actual = (document.core_properties.keywords or "").removeprefix(HASH_PREFIX)
    if actual != expected:
        raise RuntimeError(f"DOCX is stale: expected {expected}, found {actual or '<missing>'}")
    if len(document.paragraphs) < 35 or len(document.tables) < 10 or len(document.inline_shapes) < 2:
        raise RuntimeError("DOCX structural sanity check failed")
    print(f"CURRENT {OUTPUT}\nSOURCE_HASH {expected}")
    print(f"STRUCTURE paragraphs={len(document.paragraphs)} tables={len(document.tables)} images={len(document.inline_shapes)}")


def main() -> int:
    options = args()
    try:
        build() if options.build else check()
    except Exception as error:
        print(f"ERROR: {error}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
