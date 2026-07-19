#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from PIL import Image, ImageChops
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from markdown_it import MarkdownIt

import publication_v3 as pub

from design_document_diagrams import (
    build_responsibility_diagram,
    build_status_diagram,
    build_workflow_diagram,
)

ACCENT = "315EFB"
ACCENT_LIGHT = "E8EEFF"
WHITE = "FFFFFF"
LIGHT_ROW = "F4F6FA"
ALLOWED_VISUAL_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}
ACTIVE_STATUSES = {"ACTIVE", "CURRENT"}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build human-readable DOCX/PDF game design documents from structured JSON sources."
    )
    parser.add_argument("--registry", required=True, help="DESIGN_DOCUMENT_REGISTRY.json")
    parser.add_argument("--only", action="append", default=[], help="Build only the given document_id; repeatable")
    parser.add_argument("--source-commit", default=os.environ.get("GITHUB_SHA", ""))
    parser.add_argument(
        "--human-visual-review",
        default="NOT_RUN",
        choices=["NOT_RUN", "PASSED", "FAILED"],
    )
    parser.add_argument("--force", action="store_true", help="Rebuild even when the v3 input digest is unchanged")
    parser.add_argument("--preflight", action="store_true", help="Check dependencies and output paths without building")
    return parser.parse_args()


def load_json(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError(f"JSON root must be an object: {path}")
    return data


def digest(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def resolve(base: Path, value: str) -> Path:
    path = Path(value)
    return path if path.is_absolute() else (base / path).resolve()


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd")) or OxmlElement("w:shd")
    if node.getparent() is None:
        props.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if margins.getparent() is None:
        props.append(margins)
    for name, value in (("top", 80), ("start", 95), ("bottom", 80), ("end", 95)):
        node = margins.find(qn(f"w:{name}")) or OxmlElement(f"w:{name}")
        if node.getparent() is None:
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_run(run, size: float = 9, bold: bool = False, color: tuple[int, int, int] | None = None) -> None:
    run.font.name = "NanumGothic"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "NanumGothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def configure_styles(doc: Document) -> None:
    font_name = "Malgun Gothic" if os.name == "nt" else "NanumGothic"
    for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    doc.styles["Normal"].font.size = Pt(9)
    doc.styles["Title"].font.size = Pt(30)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].font.color.rgb = RGBColor(49, 94, 251)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].font.bold = True
    if "Code Block" not in [style.name for style in doc.styles]:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(8)


def _value(value: Any) -> str:
    if isinstance(value, list):
        return "\n".join(str(item) for item in value) if value else "-"
    if isinstance(value, dict):
        return " / ".join(f"{key}: {_value(item)}" for key, item in value.items()) if value else "-"
    if value in (None, ""):
        return "-"
    if isinstance(value, bool):
        return "예" if value else "아니오"
    return str(value)


def add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, ACCENT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(cell.paragraphs[0].add_run(str(header)), 7.6, True, (255, 255, 255))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else LIGHT_ROW)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            style_run(cell.paragraphs[0].add_run(_value(value)), 7.1)
    doc.add_paragraph()


def add_bullets(doc: Document, values: list[Any]) -> None:
    for value in values or ["[작성 필요]"]:
        paragraph = doc.add_paragraph(style="List Bullet")
        style_run(paragraph.add_run(_value(value)), 9)


def add_metadata_table(doc: Document, document: dict[str, Any], source_hash: str, source_commit: str) -> None:
    metadata = document.get("metadata", {})
    pairs = [
        ("프로젝트", document.get("project", "[프로젝트명]")),
        ("분야", document.get("discipline", "[분야]")),
        ("문서 종류", document.get("document_kind", "discipline-bible")),
        ("책임", document.get("owner", "[책임자]")),
        ("상태", document.get("status", "DRAFT")),
        ("마지막 검토", metadata.get("last_reviewed", "[작성 필요]")),
        ("기준 커밋", source_commit or metadata.get("source_commit", "[작성 필요]")),
        ("JSON SHA-256", source_hash[:20] + "…"),
        ("현재 제품 게이트", metadata.get("current_product_gate", "[작성 필요]")),
        ("현재 작업 게이트", metadata.get("current_work_gate", "[작성 필요]")),
        ("다음 검토 트리거", metadata.get("next_review_trigger", "[작성 필요]")),
        ("사람용 형식", "DOCX + PDF + 다이어그램·승인 이미지"),
    ]
    table = doc.add_table(rows=(len(pairs) + 1) // 2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (label, value) in enumerate(pairs):
        row, pair = divmod(index, 2)
        label_cell = table.rows[row].cells[pair * 2]
        value_cell = table.rows[row].cells[pair * 2 + 1]
        set_cell_shading(label_cell, ACCENT_LIGHT)
        set_cell_margins(label_cell)
        set_cell_margins(value_cell)
        style_run(label_cell.paragraphs[0].add_run(label), 8.5, True)
        style_run(value_cell.paragraphs[0].add_run(_value(value)), 8.5)
    doc.add_paragraph()


def add_header_footer(doc: Document, title: str, source_hash: str) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = title
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            style_run(run, 7.5, color=(102, 112, 133))
        footer = section.footer.paragraphs[0]
        footer.text = f"자동 생성 사람용 파생본 | 원본 SHA-256 {source_hash[:16]}… | "
        page_field = OxmlElement("w:fldSimple")
        page_field.set(qn("w:instr"), "PAGE")
        footer._p.append(page_field)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            style_run(run, 7, color=(102, 112, 133))


def add_toc(doc: Document, items: list[str]) -> None:
    doc.add_heading("목차", level=1)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.9) -> None:
    doc.add_picture(str(path), width=Inches(width))
    paragraph = doc.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        style_run(run, 8, color=(102, 112, 133))


def normalize_table_rows(items: list[Any], fields: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in items:
        if isinstance(item, dict):
            rows.append([_value(item.get(field)) for field in fields])
        else:
            rows.append([_value(item), *["" for _ in fields[1:]]])
    return rows


def build_docx(
    document: dict[str, Any],
    source_hash: str,
    source_commit: str,
    generated_diagrams: dict[str, Path],
    approved_visuals: list[tuple[dict[str, Any], Path]],
    output: Path,
    *,
    source_path: str,
    source_format: str,
    generated_at: str,
    epoch: int,
) -> None:
    doc = Document()
    stable_date = datetime.fromtimestamp(max(epoch, 0), timezone.utc).replace(tzinfo=None)
    doc.core_properties.created = stable_date
    doc.core_properties.modified = stable_date
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.65)
    configure_styles(doc)
    title_text = str(document.get("title", "[기획서 제목]"))
    add_header_footer(doc, title_text, source_hash)

    title = doc.add_paragraph(style="Title")
    title.add_run(title_text)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(f"{source_format.upper()} 책임 원본에서 자동 생성한 사람용 최신 통합본")
    add_metadata_table(doc, document, source_hash, source_commit)
    add_table(doc, ["발행 정보", "값"], [
        ["책임 원본 경로", source_path],
        ["원본 형식", source_format],
        ["생성 기준 시각", generated_at],
        ["편집 정책", "책임 원본을 수정한 뒤 재생성"],
    ])
    add_toc(doc, [
        "1. 한눈에 보기", "2. 목표와 Quality Bar", "3. 책임과 협업 경계",
        "4. 분야 전체 작업 과정", "5. 개발 게이트", "6. 프로젝트 스킬",
        "7. 현재 상태", "8. 확정 결정", "9. 구현·제작 상태",
        "10. 검증 상태", "승인 이미지·UI·다이어그램·실제 캡처",
        "위험·미확정·보류", "다음 작업", "Definition of Ready",
        "Definition of Done", "부록·참조", "변경·학습 이력",
    ])

    overview = document.get("overview", {})
    doc.add_heading("1. 한눈에 보기", level=1)
    add_table(doc, ["항목", "현재 기준"], [
        ["분야 목적", overview.get("purpose")],
        ["플레이어·사용자 가치", overview.get("player_value")],
        ["현재 목표", overview.get("current_goal")],
        ["현재 요약", overview.get("summary")],
        ["가장 큰 위험", overview.get("largest_risk")],
        ["다음 작업", overview.get("next_action")],
    ])
    if "status" in generated_diagrams:
        add_image(doc, generated_diagrams["status"], "그림 1. 확정·구현·검증·확인 필요·보류 상태 분리")

    doc.add_heading("2. 목표와 Quality Bar", level=1)
    add_bullets(doc, document.get("goals", []))
    quality = document.get("quality_bar", [])
    add_table(doc, ["기준", "합격 조건", "증거", "상태"], normalize_table_rows(quality, ["criterion", "pass_condition", "evidence", "status"]) or [["[작성 필요]", "-", "-", "미확정"]])
    doc.add_heading("금지 방향", level=2)
    add_bullets(doc, document.get("prohibited_directions", []))

    doc.add_heading("3. 책임과 협업 경계", level=1)
    if "responsibility" in generated_diagrams:
        add_image(doc, generated_diagrams["responsibility"], "그림 2. 분야 책임·협업 경계")
    responsibilities = document.get("responsibilities", {})
    doc.add_heading("이 분야가 소유", level=2)
    add_bullets(doc, responsibilities.get("owns", []))
    doc.add_heading("이 분야가 소유하지 않음", level=2)
    add_bullets(doc, responsibilities.get("does_not_own", []))
    interfaces = responsibilities.get("interfaces", [])
    add_table(doc, ["상대 분야", "받는 입력", "제공하는 출력", "함께 갱신할 원본"], normalize_table_rows(interfaces, ["discipline", "receives", "provides", "update_together"]) or [["[작성 필요]", "-", "-", "-"]])

    doc.add_heading("4. 분야 전체 작업 과정", level=1)
    if "workflow" in generated_diagrams:
        add_image(doc, generated_diagrams["workflow"], "그림 3. 입력부터 다음 게이트까지의 전체 작업 흐름")
    workflow = document.get("workflow", [])
    add_table(doc, ["단계", "입력", "판단·작업", "산출물", "실제 경로", "실패·폴백", "검증", "다음 게이트"], normalize_table_rows(workflow, ["step", "inputs", "decision_or_work", "outputs", "paths", "fallback", "validation", "next_gate"]) or [["[작성 필요]", "-", "-", "-", "-", "-", "-", "-"]])

    doc.add_heading("5. 개발 게이트", level=1)
    gates = document.get("development_gates", {})
    add_table(doc, ["작업 게이트", "진입 조건", "종료 기준", "증거", "상태"], normalize_table_rows(gates.get("work", []), ["gate", "entry", "exit", "evidence", "status"]) or [["[작성 필요]", "-", "-", "-", "-"]])
    add_table(doc, ["제품 단계", "이 분야가 증명할 것", "Quality Bar", "증거", "상태"], normalize_table_rows(gates.get("product", []), ["gate", "proof", "quality_bar", "evidence", "status"]) or [["[작성 필요]", "-", "-", "-", "-"]])

    doc.add_heading("6. 프로젝트 스킬", level=1)
    skills = document.get("skills", {})
    add_table(doc, ["구분", "스킬·경로"], [
        ["Foundation", skills.get("foundation")],
        ["분야 진입 스킬", skills.get("discipline_entry")],
        ["하위 스킬", skills.get("subskills")],
        ["Learning Log", skills.get("learning_log")],
        ["현재 지식 상태", skills.get("knowledge_state")],
        ["다음 검토 트리거", skills.get("next_review_trigger")],
    ])

    doc.add_heading("7. 현재 상태", level=1)
    states = document.get("current_state", [])
    add_table(doc, ["항목", "확정", "구현·제작", "검증", "확인 필요", "보류", "책임 원본·증거"], normalize_table_rows(states, ["item", "confirmed", "implemented", "validated", "unresolved", "hold", "evidence"]) or [["[작성 필요]", "-", "-", "-", "-", "-", "-"]])

    doc.add_heading("8. 확정 결정", level=1)
    add_table(doc, ["ID", "결정", "이유", "영향 범위", "검증", "관련 스킬"], normalize_table_rows(document.get("decisions", []), ["id", "decision", "reason", "impact", "validation", "skills"]) or [["[작성 필요]", "-", "-", "-", "-", "-"]])

    doc.add_heading("9. 구현·제작 상태", level=1)
    add_table(doc, ["기능·자산", "실제 경로", "반영 범위", "남은 범위", "위험", "상태"], normalize_table_rows(document.get("implementation", []), ["item", "path", "scope", "remaining", "risk", "status"]) or [["[작성 필요]", "-", "-", "-", "-", "-"]])

    doc.add_heading("10. 검증 상태", level=1)
    add_table(doc, ["대상", "종류", "명령·방법", "결과", "증거", "미검증 이유"], normalize_table_rows(document.get("validation", []), ["target", "type", "method", "result", "evidence", "not_run_reason"]) or [["[작성 필요]", "-", "-", "-", "-", "-"]])

    for section_index, custom in enumerate(document.get("sections", []), start=11):
        doc.add_heading(f"{section_index}. {custom.get('title', '[상세 기획]')}", level=1)
        purpose = custom.get("purpose")
        if purpose:
            paragraph = doc.add_paragraph()
            style_run(paragraph.add_run(str(purpose)), 9)
        for block in custom.get("blocks", []):
            block_type = block.get("type")
            if block_type == "paragraph":
                paragraph = doc.add_paragraph()
                style_run(paragraph.add_run(_value(block.get("text"))), 9)
            elif block_type == "bullets":
                add_bullets(doc, block.get("items", []))
            elif block_type == "table":
                headers = [str(value) for value in block.get("headers", [])]
                rows = [[_value(value) for value in row] for row in block.get("rows", [])]
                if headers:
                    add_table(doc, headers, rows)

    doc.add_heading("승인 이미지·UI·다이어그램·실제 캡처", level=1)
    if approved_visuals:
        for index, (item, path) in enumerate(approved_visuals, start=1):
            add_image(
                doc,
                path,
                f"승인 자료 {index}. {item.get('title', item.get('asset_id', path.name))} - "
                f"{item.get('caption', '')} | 출처: {item.get('source', '[미등록]')} | "
                f"상태: {item.get('status', '[미등록]')}",
                width=6.7,
            )
            add_table(doc, ["Asset ID", "상태", "출처", "캐노니컬 경로", "채택 요소", "비채택 요소"], [[item.get("asset_id"), item.get("status"), item.get("source"), item.get("path"), item.get("adopted_elements"), item.get("excluded_elements")]])
    else:
        doc.add_paragraph("[등록된 승인 이미지 없음]")

    doc.add_heading("위험·미확정·보류", level=1)
    add_table(doc, ["ID", "유형", "항목", "영향", "완화·재개 조건", "책임", "상태"], normalize_table_rows(document.get("risks", []), ["id", "type", "item", "impact", "mitigation_or_resume", "owner", "status"]) or [["[작성 필요]", "-", "-", "-", "-", "-", "-"]])

    doc.add_heading("다음 작업", level=1)
    add_table(doc, ["우선순위", "작업", "주 책임", "영향 분야", "선행 조건", "산출물", "완료 기준", "검증", "스킬"], normalize_table_rows(document.get("next_actions", []), ["priority", "task", "owner", "affected_disciplines", "prerequisites", "deliverables", "done", "validation", "skills"]) or [["1", "[작성 필요]", "-", "-", "-", "-", "-", "-", "-"]])

    doc.add_heading("Definition of Ready", level=1)
    add_bullets(doc, document.get("definition_of_ready", []))
    doc.add_heading("Definition of Done", level=1)
    add_bullets(doc, document.get("definition_of_done", []))

    doc.add_heading("부록·참조", level=1)
    add_table(doc, ["자료", "경로", "역할", "상태"], normalize_table_rows(document.get("appendices", []), ["title", "path", "role", "status"]) or [["[작성 필요]", "-", "-", "-"]])
    doc.add_heading("변경·학습 이력", level=1)
    add_table(doc, ["일자", "변경", "이유", "검증", "학습 상태"], normalize_table_rows(document.get("change_history", []), ["date", "change", "reason", "validation", "knowledge_state"]) or [["[작성 필요]", "-", "-", "-", "-"]])

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def parse_markdown(path: Path, entry: dict[str, Any]) -> tuple[list[Any], list[str], list[Path]]:
    source = path.read_text(encoding="utf-8")
    if re.search(r"</?[A-Za-z][^>]*>", source):
        raise ValueError(f"Raw HTML is not allowed in reproducible Markdown sources: {path}")
    parser = MarkdownIt("commonmark", {"html": False, "linkify": False}).enable("table")
    tokens = parser.parse(source)
    headings: list[tuple[int, str]] = []
    mermaid_sources: list[str] = []
    local_images: list[Path] = []
    for index, token in enumerate(tokens):
        if token.type == "heading_open" and index + 1 < len(tokens):
            headings.append((int(token.tag[1:]), tokens[index + 1].content.strip()))
        if token.type in {"html_block", "html_inline"}:
            raise ValueError(f"Raw HTML is not allowed in reproducible Markdown sources: {path}")
        if token.type == "fence" and token.info.strip().lower() == "mermaid":
            mermaid_sources.append(token.content.strip() + "\n")
        if token.type == "inline":
            for child in token.children or []:
                if child.type != "image":
                    continue
                src = child.attrGet("src") or ""
                if src.startswith(("http://", "https://", "//")):
                    raise ValueError(f"Remote images are not allowed: {src}")
                image_path = resolve(path.parent, src)
                if not image_path.is_file():
                    raise FileNotFoundError(f"Markdown image missing: {src}")
                local_images.append(image_path)
    h1 = [text for level, text in headings if level == 1]
    if len(h1) != 1 or h1[0] != entry["title"]:
        raise ValueError(f"Markdown H1 must exactly match Registry title {entry['title']!r}: {path}")
    required = entry.get("required_sections") or ["목표", "배경과 의도", "범위", "규칙과 제약", "검증과 완료 기준"]
    h2 = {text for level, text in headings if level == 2}
    missing = [section for section in required if section not in h2]
    if missing:
        raise ValueError(f"Markdown required sections missing in {path}: {', '.join(missing)}")
    if entry.get("diagram_policy") == "mermaid" and not mermaid_sources:
        raise ValueError(f"diagram_policy=mermaid requires at least one Mermaid fenced block: {path}")
    return tokens, mermaid_sources, sorted(set(local_images))


def render_mermaid(
    sources: list[str],
    output_dir: Path,
    repository_root: Path,
) -> tuple[dict[str, Path], dict[str, Path], dict[str, Path]]:
    if not sources:
        return {}, {}, {}
    executable = pub.mermaid_cli_path(repository_root)
    if not executable:
        raise RuntimeError(
            "Mermaid CLI is required. Run `pnpm install --frozen-lockfile` or set BASE_MERMAID_CLI."
        )
    chrome = pub.chrome_path()
    if not chrome:
        raise RuntimeError("Chrome/Chromium is required for Mermaid rendering; set PUPPETEER_EXECUTABLE_PATH.")
    environment = os.environ.copy()
    environment["PUPPETEER_EXECUTABLE_PATH"] = chrome
    source_paths: dict[str, Path] = {}
    svg_paths: dict[str, Path] = {}
    png_paths: dict[str, Path] = {}
    output_dir.mkdir(parents=True, exist_ok=True)
    for index, source in enumerate(sources, start=1):
        stem = f"mermaid-{index:02d}"
        source_path = output_dir / f"{stem}.mmd"
        svg_path = output_dir / f"{stem}.svg"
        png_path = output_dir / f"{stem}.png"
        source_path.write_text(source, encoding="utf-8")
        for destination in (svg_path, png_path):
            result = subprocess.run(
                [executable, "-i", str(source_path), "-o", str(destination), "--backgroundColor", "white"],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                env=environment,
                timeout=120,
            )
            if result.returncode:
                raise RuntimeError(f"Mermaid rendering failed for {source_path}:\n{result.stdout}{result.stderr}")
        source_paths[stem] = source_path
        svg_paths[stem] = svg_path
        png_paths[stem] = png_path
    return source_paths, svg_paths, png_paths


def _inline_to_paragraph(paragraph, token, source_dir: Path) -> list[tuple[Path, str]]:
    bold = False
    italic = False
    active_link = ""
    images: list[tuple[Path, str]] = []
    for child in token.children or []:
        if child.type == "strong_open":
            bold = True
        elif child.type == "strong_close":
            bold = False
        elif child.type == "em_open":
            italic = True
        elif child.type == "em_close":
            italic = False
        elif child.type == "link_open":
            active_link = child.attrGet("href") or ""
        elif child.type == "link_close":
            if active_link:
                run = paragraph.add_run(f" ({active_link})")
                style_run(run, 9, color=(49, 94, 251))
            active_link = ""
        elif child.type == "image":
            src = child.attrGet("src") or ""
            images.append((resolve(source_dir, src), child.content or Path(src).name))
        elif child.type in {"softbreak", "hardbreak"}:
            paragraph.add_run().add_break()
        elif child.type in {"text", "code_inline"}:
            run = paragraph.add_run(child.content)
            style_run(run, 8.5 if child.type == "code_inline" else 9, bold)
            run.italic = italic
            if child.type == "code_inline":
                run.font.name = "Consolas"
    return images


def _table_from_tokens(tokens: list[Any], start: int) -> tuple[list[str], list[list[str]], int]:
    headers: list[str] = []
    rows: list[list[str]] = []
    current: list[str] = []
    in_head = False
    index = start
    while index < len(tokens):
        token = tokens[index]
        if token.type == "thead_open":
            in_head = True
        elif token.type == "thead_close":
            in_head = False
        elif token.type == "tr_open":
            current = []
        elif token.type == "inline" and index > 0 and tokens[index - 1].type in {"th_open", "td_open"}:
            current.append(token.content)
        elif token.type == "tr_close":
            if in_head and not headers:
                headers = current
            else:
                rows.append(current)
        elif token.type == "table_close":
            return headers, rows, index
        index += 1
    raise ValueError("Unclosed Markdown table")


def build_markdown_docx(
    tokens: list[Any],
    entry: dict[str, Any],
    source_path: Path,
    source_hash: str,
    source_commit: str,
    generated_at: str,
    epoch: int,
    mermaid_png: dict[str, Path],
    output: Path,
) -> None:
    doc = Document()
    stable_date = datetime.fromtimestamp(max(epoch, 0), timezone.utc).replace(tzinfo=None)
    doc.core_properties.created = stable_date
    doc.core_properties.modified = stable_date
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.65)
    configure_styles(doc)
    add_header_footer(doc, entry["title"], source_hash)
    title = doc.add_paragraph(style="Title")
    title.add_run(entry["title"])
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Markdown 책임 원본에서 자동 생성한 사람용 최신 통합본")
    add_table(doc, ["발행 정보", "값"], [
        ["분야", entry["discipline"]],
        ["상태", entry["status"]],
        ["책임 원본 경로", str(entry["source_path"])],
        ["기준 커밋", source_commit],
        ["생성 기준 시각", generated_at],
        ["원본 SHA-256", source_hash],
    ])
    heading_items = [
        token.content
        for index, token in enumerate(tokens)
        if token.type == "inline" and index > 0 and tokens[index - 1].type == "heading_open"
        and tokens[index - 1].tag != "h1"
    ]
    add_toc(doc, heading_items)
    mermaid_index = 0
    list_style: str | None = None
    index = 0
    while index < len(tokens):
        token = tokens[index]
        if token.type == "heading_open":
            level = int(token.tag[1:])
            inline = tokens[index + 1]
            if level > 1:
                doc.add_heading(inline.content, level=min(level - 1, 3))
            index += 2
            continue
        if token.type == "bullet_list_open":
            list_style = "List Bullet"
        elif token.type == "ordered_list_open":
            list_style = "List Number"
        elif token.type in {"bullet_list_close", "ordered_list_close"}:
            list_style = None
        elif token.type == "paragraph_open" and index + 1 < len(tokens):
            inline = tokens[index + 1]
            paragraph = doc.add_paragraph(style=list_style if list_style else None)
            images = _inline_to_paragraph(paragraph, inline, source_path.parent)
            for image_path, caption in images:
                add_image(doc, image_path, caption, width=6.7)
            index += 2
            continue
        elif token.type == "fence":
            if token.info.strip().lower() == "mermaid":
                mermaid_index += 1
                add_image(doc, mermaid_png[f"mermaid-{mermaid_index:02d}"], f"Mermaid 다이어그램 {mermaid_index}", width=6.7)
            else:
                paragraph = doc.add_paragraph(style="Code Block")
                paragraph.add_run(token.content)
        elif token.type == "blockquote_open":
            paragraph = doc.add_paragraph("인용", style="Intense Quote")
            paragraph.runs[0].bold = True
        elif token.type == "table_open":
            headers, rows, end = _table_from_tokens(tokens, index + 1)
            add_table(doc, headers, rows)
            index = end
        index += 1
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def convert_pdf(docx_path: Path, pdf_path: Path, epoch: int = 0) -> None:
    executable = pub.libreoffice_path()
    if not executable:
        raise RuntimeError(
            "LibreOffice is required. Install it or set BASE_LIBREOFFICE; existing publications were not changed."
        )
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="design-doc-lo-") as temp:
        profile = Path(temp) / "profile"
        command = [
            executable,
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ]
        environment = os.environ.copy()
        environment["SOURCE_DATE_EPOCH"] = str(max(epoch, 0))
        try:
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=120,
                env=environment,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("LibreOffice PDF conversion timed out after 120 seconds.") from exc
        if result.returncode:
            raise RuntimeError(result.stdout + result.stderr)
    generated = pdf_path.parent / f"{docx_path.stem}.pdf"
    if generated != pdf_path:
        generated.replace(pdf_path)
    if not pdf_path.exists() or pdf_path.read_bytes()[:5] != b"%PDF-":
        raise RuntimeError(f"Invalid PDF output: {pdf_path}")


def render_pdf_for_review(pdf_path: Path, output_dir: Path) -> list[Path]:
    executable = pub.pdftoppm_path()
    if not executable:
        raise RuntimeError(
            "pdftoppm is required. Install Poppler or set BASE_PDFTOPPM; existing publications were not changed."
        )
    output_dir.mkdir(parents=True, exist_ok=True)
    prefix = output_dir / "page"
    result = subprocess.run(
        [executable, "-png", "-r", "120", str(pdf_path), str(prefix)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=120,
    )
    if result.returncode:
        raise RuntimeError(result.stdout + result.stderr)
    pages = sorted(output_dir.glob("page-*.png"))
    if not pages:
        raise RuntimeError("PDF render review produced no pages.")
    for page in pages:
        with Image.open(page).convert("RGB") as image:
            if image.width < 500 or image.height < 500:
                raise RuntimeError(f"Rendered page is unexpectedly small: {page}")
            white = Image.new("RGB", image.size, "white")
            if ImageChops.difference(image, white).getbbox() is None:
                raise RuntimeError(f"Rendered page is blank: {page}")
    return pages


def approved_visuals(document: dict[str, Any], registry_dir: Path) -> list[tuple[dict[str, Any], Path]]:
    result: list[tuple[dict[str, Any], Path]] = []
    for item in document.get("approved_visuals", []):
        if not item.get("include_in_publication", True):
            continue
        path_value = str(item.get("path", "")).strip()
        if not path_value:
            continue
        path = resolve(registry_dir, path_value)
        if not path.is_file():
            raise FileNotFoundError(f"Approved visual missing: {path_value}")
        if path.suffix.lower() not in ALLOWED_VISUAL_SUFFIXES:
            raise ValueError(f"Unsupported approved visual format: {path_value}")
        result.append((item, path))
    return result


def _text_digest(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _current_commit() -> str:
    result = subprocess.run(
        ["git", "rev-parse", "HEAD"],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        check=False,
    )
    return result.stdout.strip() if result.returncode == 0 else ""


def _existing_current(manifest_path: Path, input_hash: str, output_pdf: Path, output_docx: Path | None) -> dict[str, Any] | None:
    if not manifest_path.is_file() or not output_pdf.is_file():
        return None
    try:
        manifest = pub.load_json(manifest_path)
        pub.require_schema_v3(manifest, str(manifest_path))
    except (ValueError, json.JSONDecodeError):
        return None
    if manifest.get("input_sha256") != input_hash or manifest.get("sync_status") != "CURRENT":
        return None
    if manifest.get("output_pdf_sha256") != pub.digest(output_pdf):
        return None
    if output_docx is not None:
        if not output_docx.is_file() or manifest.get("output_docx_sha256") != pub.digest(output_docx):
            return None
    return manifest


def build_one(
    entry: dict[str, Any],
    registry_path: Path,
    source_commit: str,
    human_review: str,
    *,
    force: bool,
) -> dict[str, Any]:
    registry_dir = registry_path.parent
    repository_root = Path(__file__).resolve().parents[1]
    schema_dir = repository_root / "schemas"
    source_path = pub.resolve(registry_dir, entry["source_path"])
    output_pdf = pub.resolve(registry_dir, entry["output_pdf"])
    output_docx = pub.resolve(registry_dir, entry["output_docx"]) if entry.get("output_docx") else None
    asset_dir = pub.resolve(registry_dir, entry["asset_dir"]) if entry.get("asset_dir") else None
    manifest_path = pub.resolve(registry_dir, entry["publication_manifest"])
    generator_path = Path(__file__).resolve()
    diagram_module = generator_path.with_name("design_document_diagrams.py")
    source_format = entry["source_format"]
    source_hash = pub.digest(source_path)
    document: dict[str, Any] | None = None
    tokens: list[Any] = []
    mermaid_sources: list[str] = []
    local_images: list[Path] = []
    visuals: list[tuple[dict[str, Any], Path]] = []
    if source_format == "json":
        document = pub.load_json(source_path)
        pub.require_schema_v3(document, str(source_path))
        pub.validate_json(document, schema_dir / "structured-design-document-v3.schema.json", str(source_path))
        if document.get("document_id") != entry["document_id"]:
            raise ValueError(f"document_id mismatch: {source_path}")
        if document.get("title") != entry["title"]:
            raise ValueError(f"title mismatch between Registry and source: {source_path}")
        visuals = approved_visuals(document, registry_dir)
    elif source_format == "markdown":
        tokens, mermaid_sources, local_images = parse_markdown(source_path, entry)
    else:
        raise ValueError(f"Unsupported source_format: {source_format}")

    input_items = [
        ("source", source_hash),
        ("generator", pub.digest(generator_path)),
        ("registry-entry", _text_digest(json.dumps(entry, ensure_ascii=False, sort_keys=True))),
        ("registry-schema", pub.digest(schema_dir / "design-document-registry-v3.schema.json")),
        ("manifest-schema", pub.digest(schema_dir / "publication-manifest-v3.schema.json")),
        ("source-commit", _text_digest(source_commit)),
    ]
    if entry["diagram_policy"] == "generated":
        input_items.append(("diagram-generator", pub.digest(diagram_module)))
    for _, path in visuals:
        input_items.append((f"approved:{os.path.relpath(path, registry_dir)}", pub.digest(path)))
    for path in local_images:
        input_items.append((f"image:{os.path.relpath(path, source_path.parent)}", pub.digest(path)))
    if mermaid_sources:
        input_items.append(("mermaid-lock", pub.digest(repository_root / "pnpm-lock.yaml")))
    input_hash = pub.combined_digest(input_items)
    if not force:
        existing = _existing_current(manifest_path, input_hash, output_pdf, output_docx)
        if existing:
            return existing

    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    epoch = pub.source_epoch(source_commit)
    generated_at = pub.stable_timestamp(epoch)
    with tempfile.TemporaryDirectory(prefix=f"{entry['document_id']}-v3-", dir=manifest_path.parent) as temp_name:
        temp_root = Path(temp_name)
        temp_assets = temp_root / "assets" / "generated"
        temp_assets.mkdir(parents=True)
        generated_assets: dict[str, Path] = {}
        mermaid_mmd: dict[str, Path] = {}
        mermaid_svg: dict[str, Path] = {}
        mermaid_png: dict[str, Path] = {}
        if entry["diagram_policy"] == "generated":
            if document is None or asset_dir is None:
                raise ValueError("diagram_policy=generated requires a JSON source and asset_dir")
            generated_assets = {
                "workflow": temp_assets / "workflow.png",
                "status": temp_assets / "status-summary.png",
                "responsibility": temp_assets / "responsibility-map.png",
            }
            build_workflow_diagram(document, generated_assets["workflow"])
            build_status_diagram(document, generated_assets["status"])
            build_responsibility_diagram(document, generated_assets["responsibility"])
        elif entry["diagram_policy"] == "mermaid":
            if asset_dir is None:
                raise ValueError("diagram_policy=mermaid requires asset_dir")
            mermaid_mmd, mermaid_svg, mermaid_png = render_mermaid(mermaid_sources, temp_assets, repository_root)

        temp_docx = temp_root / f"{output_pdf.stem}.docx"
        temp_pdf = temp_root / output_pdf.name
        if source_format == "json":
            assert document is not None
            build_docx(
                document,
                source_hash,
                source_commit,
                generated_assets,
                visuals,
                temp_docx,
                source_path=entry["source_path"],
                source_format=source_format,
                generated_at=generated_at,
                epoch=epoch,
            )
        else:
            build_markdown_docx(
                tokens,
                entry,
                source_path,
                source_hash,
                source_commit,
                generated_at,
                epoch,
                mermaid_png,
                temp_docx,
            )
        pub.normalize_docx(temp_docx, epoch)
        convert_pdf(temp_docx, temp_pdf, epoch)
        pub.normalize_pdf(temp_pdf, epoch, input_hash)
        with tempfile.TemporaryDirectory(prefix="render-review-") as review_name:
            page_count = len(render_pdf_for_review(temp_pdf, Path(review_name)))

        pdf_hash = pub.digest(temp_pdf)
        docx_hash = pub.digest(temp_docx) if output_docx is not None else None
        target_assets: dict[Path, Path] = {}
        if asset_dir is not None:
            for path in generated_assets.values():
                target_assets[path] = asset_dir / "generated" / path.name
            for collection in (mermaid_mmd, mermaid_svg, mermaid_png):
                for path in collection.values():
                    target_assets[path] = asset_dir / "generated" / path.name

        def asset_hashes(collection: dict[str, Path], suffix: str) -> dict[str, str]:
            if asset_dir is None:
                return {}
            return {
                os.path.relpath(asset_dir / "generated" / f"{name}.{suffix}", registry_dir): pub.digest(path)
                for name, path in collection.items()
            }

        generated_hashes = {
            os.path.relpath(target, registry_dir): pub.digest(source)
            for source, target in target_assets.items()
        }
        manifest = {
            "schema_version": 3,
            "publication_id": entry["document_id"],
            "role": "human-readable-derivative",
            "source_path": entry["source_path"],
            "source_format": source_format,
            "source_sha256": source_hash,
            "input_sha256": input_hash,
            "source_commit": source_commit,
            "generator": entry["generator"],
            "generator_sha256": pub.digest(generator_path),
            "generated_at": generated_at,
            "output_docx": entry.get("output_docx"),
            "output_docx_sha256": docx_hash,
            "output_pdf": entry["output_pdf"],
            "output_pdf_sha256": pdf_hash,
            "asset_dir": entry.get("asset_dir"),
            "generated_assets": generated_hashes,
            "approved_visuals": {os.path.relpath(path, registry_dir): pub.digest(path) for _, path in visuals},
            "source_images": {os.path.relpath(path, registry_dir): pub.digest(path) for path in local_images},
            "mermaid_sources": asset_hashes(mermaid_mmd, "mmd"),
            "mermaid_svg": asset_hashes(mermaid_svg, "svg"),
            "mermaid_png": asset_hashes(mermaid_png, "png"),
            "sync_status": "CURRENT",
            "automated_render_review": "PASSED",
            "rendered_page_count": page_count,
            "human_visual_review": human_review,
            "human_visual_review_pdf_sha256": pdf_hash if human_review == "PASSED" else None,
            "editing_policy": "Edit the declared Markdown or JSON source and regenerate; do not edit derivatives.",
        }
        pub.validate_json(manifest, schema_dir / "publication-manifest-v3.schema.json", str(manifest_path))
        temp_manifest = temp_root / manifest_path.name
        pub.write_json(temp_manifest, manifest)

        for source, target in target_assets.items():
            pub.atomic_replace(source, target)
        if output_docx is not None:
            pub.atomic_replace(temp_docx, output_docx)
        pub.atomic_replace(temp_pdf, output_pdf)
        pub.atomic_replace(temp_manifest, manifest_path)
        return manifest


def preflight(registry_path: Path, registry: dict[str, Any]) -> int:
    repository_root = Path(__file__).resolve().parents[1]
    active = [item for item in registry.get("documents", []) if item.get("status") in pub.ACTIVE_STATUSES]
    checks: dict[str, Any] = {
        "python": shutil.which("python") or shutil.which("python3") or os.sys.executable,
        "libreoffice": pub.libreoffice_path(),
        "pdftoppm": pub.pdftoppm_path(),
        "font_regular": pub.font_paths()[0],
        "font_bold": pub.font_paths()[1],
        "mermaid_cli": pub.mermaid_cli_path(repository_root) if any(item.get("diagram_policy") == "mermaid" for item in active) else "NOT_REQUIRED",
        "chrome": pub.chrome_path() if any(item.get("diagram_policy") == "mermaid" for item in active) else "NOT_REQUIRED",
        "sources": {},
        "output_writable": {},
    }
    registry_dir = registry_path.parent
    for item in active:
        source = pub.resolve(registry_dir, item["source_path"])
        output = pub.resolve(registry_dir, item["output_pdf"])
        checks["sources"][item["document_id"]] = str(source) if source.is_file() else None
        output.parent.mkdir(parents=True, exist_ok=True)
        checks["output_writable"][item["document_id"]] = os.access(output.parent, os.W_OK)
    required_values = [checks["libreoffice"], checks["pdftoppm"]]
    if any(item.get("diagram_policy") in {"generated", "mermaid"} for item in active):
        required_values.extend([checks["font_regular"], checks["font_bold"]])
    if any(item.get("diagram_policy") == "mermaid" for item in active):
        required_values.extend([checks["mermaid_cli"], checks["chrome"]])
    success = all(required_values) and all(checks["sources"].values()) and all(checks["output_writable"].values())
    checks["status"] = "PASSED" if success else "FAILED"
    print(json.dumps(checks, ensure_ascii=False, indent=2))
    if not success:
        print(
            "Install the reported dependency or set BASE_LIBREOFFICE, BASE_PDFTOPPM, "
            "BASE_MERMAID_CLI, BASE_FONT_REGULAR/BASE_FONT_BOLD, or PUPPETEER_EXECUTABLE_PATH; then rerun --preflight."
        )
    return 0 if success else 1


def main() -> int:
    options = parse_args()
    registry_path = Path(options.registry).resolve()
    registry = pub.load_json(registry_path)
    pub.require_schema_v3(registry, str(registry_path))
    schema_path = Path(__file__).resolve().parents[1] / "schemas" / "design-document-registry-v3.schema.json"
    pub.validate_json(registry, schema_path, str(registry_path))
    if options.preflight:
        return preflight(registry_path, registry)
    documents = registry.get("documents", [])
    selected = set(options.only)
    source_commit = options.source_commit or _current_commit()
    built = []
    for entry in documents:
        if entry.get("status") not in pub.ACTIVE_STATUSES:
            continue
        if selected and entry.get("document_id") not in selected:
            continue
        built.append(
            build_one(
                entry,
                registry_path,
                source_commit,
                options.human_visual_review,
                force=options.force,
            )
        )
    if selected and len(built) != len(selected):
        built_ids = {item["publication_id"] for item in built}
        missing = sorted(selected - built_ids)
        raise SystemExit(f"Requested active document_id not found: {', '.join(missing)}")
    print(f"Generated or confirmed {len(built)} design document publication(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
