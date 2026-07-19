#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

import publication_v3 as pub
from skill_map_diagrams import DISCIPLINES, LAYER_KO, STATUS_KO, discipline_diagram, flow_diagram, matrix_diagram

ACCENT, ACCENT_LIGHT, WHITE = "315EFB", "E8EEFF", "FFFFFF"


def args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate human-readable DOCX/PDF project skill maps from SKILL_REGISTRY.json")
    parser.add_argument("--registry", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--project-name", default="[프로젝트명]")
    parser.add_argument("--source-commit", default=os.environ.get("GITHUB_SHA", ""))
    parser.add_argument("--human-visual-review", default="NOT_RUN", choices=["NOT_RUN", "PASSED", "FAILED"])
    parser.add_argument("--force", action="store_true")
    parser.add_argument("--preflight", action="store_true")
    return parser.parse_args()


def digest(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def load(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Registry root must be an object")
    data.setdefault("routing_policy", {})
    data.setdefault("skills", [])
    data.setdefault("discipline_entrypoints", {})
    return data


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd")) or OxmlElement("w:shd")
    if node.getparent() is None:
        props.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell) -> None:
    props = cell._tc.get_or_add_tcPr()
    mar = props.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if mar.getparent() is None:
        props.append(mar)
    for name, value in (("top", 80), ("start", 100), ("bottom", 80), ("end", 100)):
        node = mar.find(qn(f"w:{name}")) or OxmlElement(f"w:{name}")
        if node.getparent() is None:
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def run_font(run, size=8, bold=False, color=None) -> None:
    font_name = "Malgun Gothic" if os.name == "nt" else "NanumGothic"
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = result.rows[0].cells[index]
        shade(cell, ACCENT)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_font(cell.paragraphs[0].add_run(header), 8, True, (255, 255, 255))
    for row_index, values in enumerate(rows):
        cells = result.add_row().cells
        for index, value in enumerate(values):
            shade(cells[index], WHITE if row_index % 2 == 0 else "F4F6FA")
            margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run_font(cells[index].paragraphs[0].add_run(value), 7.5)
    doc.add_paragraph()


def set_styles(doc: Document) -> None:
    font_name = "Malgun Gothic" if os.name == "nt" else "NanumGothic"
    for name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = font_name
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    doc.styles["Normal"].font.size = Pt(9)
    doc.styles["Title"].font.size = Pt(30)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].font.color.rgb = RGBColor(49, 94, 251)


def header_footer(doc: Document, project_name: str, registry_hash: str) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = f"{project_name} | PROJECT SKILL MAP"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run_font(run, 8, color=(102, 112, 133))
        footer = section.footer.paragraphs[0]
        footer.text = f"Registry SHA-256: {registry_hash[:16]}… | 자동 생성 파생본 | "
        page_field = OxmlElement("w:fldSimple")
        page_field.set(qn("w:instr"), "PAGE")
        footer._p.append(page_field)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run_font(run, 7.5, color=(102, 112, 133))


def build_docx(registry: dict[str, Any], project_name: str, registry_hash: str, source_commit: str, assets: dict[str, Path], output: Path, epoch: int) -> None:
    doc = Document()
    stable_date = datetime.fromtimestamp(max(epoch, 0), timezone.utc).replace(tzinfo=None)
    doc.core_properties.created = stable_date
    doc.core_properties.modified = stable_date
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.7)
    set_styles(doc)
    header_footer(doc, project_name, registry_hash)

    title = doc.add_paragraph(style="Title")
    title.add_run(f"{project_name}\n프로젝트 스킬 지도")
    subtitle = doc.add_paragraph()
    run_font(subtitle.add_run("사람용 통합 열람본 · 이미지 포함 · SKILL_REGISTRY.json에서 자동 생성"), 11, color=(102, 112, 133))

    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    pairs = [("책임 원본", "SKILL_REGISTRY.json"), ("생성 역할", "읽기 전용 파생본"), ("Registry 해시", registry_hash[:20] + "…"), ("기준 커밋", source_commit or "[작성 필요]")]
    for index, (label, value) in enumerate(pairs):
        row, pair = divmod(index, 2)
        label_cell, value_cell = meta.rows[row].cells[pair * 2], meta.rows[row].cells[pair * 2 + 1]
        shade(label_cell, ACCENT_LIGHT)
        margins(label_cell); margins(value_cell)
        run_font(label_cell.paragraphs[0].add_run(label), 9, True)
        run_font(value_cell.paragraphs[0].add_run(value), 9)
    doc.add_paragraph()
    doc.add_heading("목차", level=1)
    for item in [
        "1. 한눈에 보기", "2. 분야별 진입 구조", "3. 현행 스킬 매트릭스",
        "4. 선택적 호출 규칙", "5. 학습·갱신 계약", "6. 운영·검수 체크",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    policy = registry["routing_policy"]
    active = [item for item in registry["skills"] if item.get("status") in ("ACTIVE", "SUPPORT")]
    coverage = sum(1 for discipline in DISCIPLINES if registry["discipline_entrypoints"].get(discipline))
    doc.add_heading("1. 한눈에 보기", level=1)
    table(doc, ["항목", "현재 기준"], [
        ["전체 스킬 자동 로드", "금지" if not policy.get("load_all_skills", False) else "허용"],
        ["기본 선택", str(policy.get("default_selection", "none"))], ["현행·보조 스킬", str(len(active))],
        ["분야 진입 스킬 등록", f"{coverage}/{len(DISCIPLINES)}"], ["주 책임 분야 스킬 최대", str(policy.get("max_primary_discipline_skills", 1))],
        ["Foundation 스킬 최대", str(policy.get("max_foundation_skills", 3))],
    ])
    doc.add_picture(str(assets["flow"]), width=Inches(7.1))
    caption = doc.add_paragraph("그림 1. 요청부터 Learning Log까지의 선택적 호출 흐름"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("2. 분야별 진입 구조", level=1)
    doc.add_picture(str(assets["discipline"]), width=Inches(7.1))
    caption = doc.add_paragraph("그림 2. 프로젝트의 11개 필수 책임 분야 진입 스킬 등록 상태"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table(doc, ["분야", "진입 스킬", "상태"], [[discipline, ", ".join(registry["discipline_entrypoints"].get(discipline, [])) or "[설치 필요]", "등록" if registry["discipline_entrypoints"].get(discipline) else "미등록"] for discipline in DISCIPLINES])

    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = Cm(1.4)
    section.left_margin = section.right_margin = Cm(1.2)
    doc.add_heading("3. 현행 스킬 매트릭스", level=1)
    doc.add_picture(str(assets["matrix"]), width=Inches(9.7))
    caption = doc.add_paragraph("그림 3. Registry 기반 스킬 상태와 호출 조건 요약"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    skills = registry["skills"] or [registry.get("skill_contract_example", {})]
    table(doc, ["스킬", "계층", "분야", "상태", "Trigger", "사용 조건", "비호출 조건", "지식 상태"], [[
        str(item.get("skill_id", "[스킬 등록 필요]")), LAYER_KO.get(str(item.get("layer", "")), str(item.get("layer", ""))), str(item.get("discipline", "-")),
        STATUS_KO.get(str(item.get("status", "NOT_INSTALLED")), str(item.get("status", ""))), " / ".join(item.get("trigger_tags", [])) or "-",
        "\n".join(item.get("use_when", [])[:2]) or "-", "\n".join(item.get("do_not_use_when", [])[:2]) or "-", str(item.get("knowledge_state", "OBSERVATION")),
    ] for item in skills])

    doc.add_heading("4. 선택적 호출 규칙", level=1)
    for text in ["전체 skills 폴더를 기본 컨텍스트로 읽지 않는다.", "trigger가 일치하지 않는 스킬은 호출하지 않는다.", "주 책임 분야 스킬은 동시에 최대 1개만 선택한다.", "Foundation 스킬은 현재 절차에 필요한 최소 개수만 선택한다.", "PDF 발행·Handoff·Health Review는 해당 단계에서만 호출한다.", "보류·백업·제거 후보 스킬은 재개 승인 전 호출하지 않는다."]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("5. 학습·갱신 계약", level=1)
    table(doc, ["대상", "계약"], [
        ["기록 조건", "실패·중요 결정·재사용 가능한 교훈·실제 검증 결과를 Learning Log에 기록"],
        ["스킬 본문 갱신", "반복 실패, 새 예외, 책임·경로·검증 변경처럼 근거가 있을 때만 수행"],
        ["지식 상태", "관찰 → 가설 → 패턴 → 검증 → 승격 후보"], ["공용화 경계", "프로젝트 고유 수치·세계관·승인 자산은 프로젝트에 유지"],
    ])
    doc.add_heading("6. 운영·검수 체크", level=1)
    for text in ["Registry의 현행 스킬 경로가 실제 존재한다.", "프로젝트의 11개 필수 분야마다 독립 진입 스킬이 등록돼 있다.", "PDF와 설정한 Markdown·DOCX·다이어그램은 Registry 해시와 일치한다.", "사람은 이 PDF를 보고, AI는 Registry를 책임 원본으로 읽는다.", "스킬 변경 시 Registry·Learning Log·PDF와 설정한 선택 파생본을 같은 작업에서 갱신한다.", "문서가 생성됐다는 사실만으로 실제 스킬 실행·검증 완료로 표시하지 않는다."]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(output)


def convert_pdf(docx: Path, output_dir: Path, epoch: int) -> Path:
    executable = pub.libreoffice_path()
    if not executable:
        raise RuntimeError("LibreOffice is required to generate PROJECT_SKILL_MAP.pdf")
    with tempfile.TemporaryDirectory(prefix="skill-map-lo-") as temp:
        profile = Path(temp) / "profile"
        profile.mkdir()
        command = [
            executable,
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx),
        ]
        try:
            environment = os.environ.copy()
            environment["SOURCE_DATE_EPOCH"] = str(max(epoch, 0))
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=120,
                env=environment,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("LibreOffice skill-map conversion timed out after 120 seconds.") from exc
        if result.returncode:
            raise RuntimeError(result.stdout + result.stderr)
    pdf = output_dir / f"{docx.stem}.pdf"
    if not pdf.exists() or pdf.read_bytes()[:5] != b"%PDF-":
        raise RuntimeError("LibreOffice did not create a valid PDF")
    return pdf


def markdown_summary(registry: dict[str, Any], registry_hash: str) -> str:
    lines = [
        "# 프로젝트 스킬 지도",
        "",
        "> 자동 생성 파생본입니다. 수동 편집하지 마세요.",
        f"> SKILL_REGISTRY.json SHA-256: `{registry_hash}`",
        "",
        "## 필수 분야",
        "",
    ]
    selected = registry.get("required_disciplines", [])
    lines.extend(f"- {item}" for item in selected) if selected else lines.append("- 필수 분야 누락")
    lines.extend(["", "## 분야별 진입 스킬", "", "| 분야 | 진입 스킬 |", "|---|---|"])
    for discipline in selected:
        entrypoints = registry.get("discipline_entrypoints", {}).get(discipline, [])
        lines.append(f"| {discipline} | {', '.join(entrypoints) or '[진입 스킬 필요]'} |")
    lines.extend(["", "## 활성 스킬", "", "| Skill ID | 계층 | 분야 | Trigger |", "|---|---|---|---|"])
    for skill in registry.get("skills", []):
        if skill.get("status") not in {"ACTIVE", "SUPPORT"}:
            continue
        lines.append(
            f"| {skill.get('skill_id', '')} | {skill.get('layer', '')} | {skill.get('discipline', '')} | "
            f"{', '.join(skill.get('trigger_tags', []))} |"
        )
    return "\n".join(lines) + "\n"


def _current_commit() -> str:
    result = subprocess.run(["git", "rev-parse", "HEAD"], capture_output=True, text=True, encoding="utf-8", errors="replace")
    return result.stdout.strip() if result.returncode == 0 else ""


def main() -> int:
    options = args()
    registry_path, output_dir = Path(options.registry).resolve(), Path(options.output_dir).resolve()
    repository_root = Path(__file__).resolve().parents[1]
    schema_dir = repository_root / "schemas"
    registry = pub.load_json(registry_path)
    pub.require_schema_v3(registry, str(registry_path))
    pub.validate_json(registry, schema_dir / "skill-registry-v3.schema.json", str(registry_path))
    if options.preflight:
        regular, bold = pub.font_paths()
        state = {
            "libreoffice": pub.libreoffice_path(),
            "pdftoppm": pub.pdftoppm_path(),
            "font_regular": regular,
            "font_bold": bold,
            "output_writable": os.access(output_dir if output_dir.exists() else output_dir.parent, os.W_OK),
        }
        state["status"] = "PASSED" if all(state.values()) else "FAILED"
        print(json.dumps(state, ensure_ascii=False, indent=2))
        return 0 if state["status"] == "PASSED" else 1

    output_dir.mkdir(parents=True, exist_ok=True)
    human = registry["human_presentation"]
    source_commit = options.source_commit or _current_commit()
    epoch = pub.source_epoch(source_commit)
    generated_at = pub.stable_timestamp(epoch)
    registry_hash = pub.digest(registry_path)
    generator_path = Path(__file__).resolve()
    diagram_path = generator_path.with_name("skill_map_diagrams.py")
    input_hash = pub.combined_digest(
        [
            ("registry", registry_hash),
            ("generator", pub.digest(generator_path)),
            ("diagram-generator", pub.digest(diagram_path)),
            ("schema", pub.digest(schema_dir / "skill-registry-v3.schema.json")),
            ("manifest-schema", pub.digest(schema_dir / "publication-manifest-v3.schema.json")),
            ("source-commit", hashlib.sha256(source_commit.encode("utf-8")).hexdigest()),
        ]
    )
    pdf_target = output_dir / human["primary_reading_format"]
    docx_target = output_dir / human["editable_derivative"] if human.get("editable_derivative") else None
    markdown_target = output_dir / human["markdown_summary"] if human.get("markdown_summary") else None
    manifest_target = output_dir / human["publication_manifest"]
    if not options.force and manifest_target.is_file() and pdf_target.is_file():
        existing = pub.load_json(manifest_target)
        if (
            existing.get("schema_version") == 3
            and existing.get("input_sha256") == input_hash
            and existing.get("sync_status") == "CURRENT"
            and existing.get("output_pdf_sha256") == pub.digest(pdf_target)
            and (docx_target is None or (docx_target.is_file() and existing.get("output_docx_sha256") == pub.digest(docx_target)))
            and (markdown_target is None or (markdown_target.is_file() and existing.get("markdown_summary_sha256") == pub.digest(markdown_target)))
        ):
            print("Skill-map publication is already CURRENT; no files rewritten.")
            return 0

    with tempfile.TemporaryDirectory(prefix="skill-map-v3-", dir=output_dir) as temp_name:
        temp_root = Path(temp_name)
        assets_dir = temp_root / "PROJECT_SKILL_MAP.assets"
        assets_dir.mkdir()
        assets = {
            "flow": assets_dir / "skill-flow.png",
            "discipline": assets_dir / "discipline-routing.png",
            "matrix": assets_dir / "skill-matrix.png",
        }
        flow_diagram(assets["flow"])
        discipline_diagram(registry, assets["discipline"])
        matrix_diagram(registry, assets["matrix"])
        temp_docx = temp_root / "PROJECT_SKILL_MAP.docx"
        build_docx(registry, options.project_name, registry_hash, source_commit, assets, temp_docx, epoch)
        pub.normalize_docx(temp_docx, epoch)
        temp_pdf = convert_pdf(temp_docx, temp_root, epoch)
        pub.normalize_pdf(temp_pdf, epoch, input_hash)
        with tempfile.TemporaryDirectory(prefix="skill-map-render-") as render_name:
            page_count = len(pub.render_pdf_for_review(temp_pdf, Path(render_name)))
        temp_markdown = temp_root / "PROJECT_SKILL_MAP.md"
        if markdown_target is not None:
            temp_markdown.write_text(markdown_summary(registry, registry_hash), encoding="utf-8")
        pdf_hash = pub.digest(temp_pdf)
        manifest = {
            "schema_version": 3,
            "publication_id": "project-skill-map",
            "role": "human-readable-derivative",
            "source_path": registry_path.name if registry_path.parent == output_dir else str(registry_path),
            "source_format": "skill-registry",
            "source_sha256": registry_hash,
            "input_sha256": input_hash,
            "source_commit": source_commit,
            "generator": "tools/build_project_skill_map.py",
            "generator_sha256": pub.digest(generator_path),
            "generated_at": generated_at,
            "output_docx": human.get("editable_derivative"),
            "output_docx_sha256": pub.digest(temp_docx) if docx_target is not None else None,
            "output_pdf": human["primary_reading_format"],
            "output_pdf_sha256": pdf_hash,
            "generated_assets": {str(Path(human.get("diagram_directory") or "PROJECT_SKILL_MAP.assets") / path.name): pub.digest(path) for path in assets.values()},
            "approved_visuals": {},
            "mermaid_sources": {},
            "mermaid_svg": {},
            "mermaid_png": {},
            "markdown_summary": human.get("markdown_summary"),
            "markdown_summary_sha256": pub.digest(temp_markdown) if markdown_target is not None else None,
            "sync_status": "CURRENT",
            "automated_render_review": "PASSED",
            "human_visual_review": options.human_visual_review,
            "human_visual_review_pdf_sha256": pdf_hash if options.human_visual_review == "PASSED" else None,
            "rendered_page_count": page_count,
            "editing_policy": "Edit SKILL_REGISTRY.json and regenerate; generated Markdown, DOCX, PDF and diagrams are read-only derivatives.",
        }
        pub.validate_json(manifest, schema_dir / "publication-manifest-v3.schema.json", str(manifest_target))
        temp_manifest = temp_root / manifest_target.name
        pub.write_json(temp_manifest, manifest)
        final_asset_dir = output_dir / (human.get("diagram_directory") or "PROJECT_SKILL_MAP.assets")
        for path in assets.values():
            pub.atomic_replace(path, final_asset_dir / path.name)
        if docx_target is not None:
            pub.atomic_replace(temp_docx, docx_target)
        if markdown_target is not None:
            pub.atomic_replace(temp_markdown, markdown_target)
        pub.atomic_replace(temp_pdf, pdf_target)
        pub.atomic_replace(temp_manifest, manifest_target)
    print(f"Generated {pdf_target}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
